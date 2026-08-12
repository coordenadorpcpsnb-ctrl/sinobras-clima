#!/usr/bin/env python3
"""
gerar_relatorio.py — Relatório executivo (2 páginas) do clima operacional
Sinobras Florestal · executado pelo GitHub Actions após update_dashboard.py

Lê os dados produzidos pela execução do dia e monta o .docx:
  data/scen12_best.csv     → totais, anomalia e probabilidades por cenário
  data/bh_final.json       → déficit, excedente e ARM mensal por cenário
  data/fc_results_best.json→ horizonte e precipitação mensal projetada
  data/serie_subst.csv     → média histórica de referência
  docs/index.html          → bloco CPC_IRI (curado manualmente a cada emissão)

Saída: docs/relatorio-executivo.docx
"""

import re, json, sys
from pathlib import Path
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = Path(__file__).parent.parent
DATA   = ROOT / 'data'
DOCS   = ROOT / 'docs'
SAIDA  = DOCS / 'relatorio-executivo.docx'

VERDE   = RGBColor(0x1F, 0x4D, 0x2E)
VERDE_C = RGBColor(0x3E, 0x7D, 0x52)
SECA    = RGBColor(0xA8, 0x32, 0x2C)
CINZA   = RGBColor(0x59, 0x59, 0x59)
PRETO   = RGBColor(0x00, 0x00, 0x00)

VERDE_HEX = '1F4D2E'
ZEBRA_HEX = 'F2F6F3'
CLIM_HEX  = 'E8EEE9'

MESES = ['jan','fev','mar','abr','mai','jun','jul','ago','set','out','nov','dez']
MESES_EXT = ['janeiro','fevereiro','março','abril','maio','junho',
             'julho','agosto','setembro','outubro','novembro','dezembro']
MESES_CAP = ['Jan','Fev','Mar','Abr','Mai','Jun','Jul','Ago','Set','Out','Nov','Dez']
ARM_CRITICO = 20   # mm — solo considerado esgotado abaixo disso


# ══════════════════════════════════════════════════════════════════════════
# Helpers de formatação XML (python-docx não expõe sombreamento e bordas)
# ══════════════════════════════════════════════════════════════════════════
def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def borda_inferior(par, cor='3E7D52', tamanho=6):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(tamanho))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), cor)
    pbdr.append(bottom)
    pPr.append(pbdr)


def borda_superior(par, cor='BFBFBF', tamanho=6):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(tamanho))
    top.set(qn('w:space'), '6')
    top.set(qn('w:color'), cor)
    pbdr.append(top)
    pPr.append(pbdr)


def bordas_tabela(tabela):
    """Linhas horizontais discretas, sem bordas verticais."""
    tblPr = tabela._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for lado, cor, sz in [('top','BFBFBF',4), ('bottom','BFBFBF',4),
                           ('left','FFFFFF',0), ('right','FFFFFF',0),
                           ('insideH','D9D9D9',2), ('insideV','FFFFFF',0)]:
        el = OxmlElement(f'w:{lado}')
        el.set(qn('w:val'), 'none' if sz == 0 else 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), cor)
        borders.append(el)
    tblPr.append(borders)


def layout_fixo(tabela, larguras):
    """python-docx ignora a largura das células sem layout fixo + tblGrid."""
    tabela.autofit = False
    tblPr = tabela._tbl.tblPr
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tblPr.append(lay)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), str(sum(larguras)))
    w.set(qn('w:type'), 'dxa')
    tblPr.append(w)
    # o tblGrid é quem manda no Word — sem ele as larguras de célula são ignoradas
    grid = tabela._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for gc, lg in zip(grid.findall(qn('w:gridCol')), larguras):
            gc.set(qn('w:w'), str(lg))
    for linha in tabela.rows:
        for c, lg in zip(linha.cells, larguras):
            c.width = Twips(lg)


def espacamento_celula(tabela, cima=60, baixo=60, esq=90, dir_=90):
    tblPr = tabela._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for lado, v in [('top',cima), ('bottom',baixo), ('left',esq), ('right',dir_)]:
        el = OxmlElement(f'w:{lado}')
        el.set(qn('w:w'), str(v))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)


def run(par, texto, *, tam=9.5, bold=False, italic=False, cor=PRETO, espaco=None):
    r = par.add_run(texto)
    r.font.name = 'Calibri'
    r.font.size = Pt(tam)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = cor
    if espaco:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(espaco))
        rPr.append(sp)
    return r


def par_corpo(doc, justificado=True, depois=5, antes=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(depois)
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.line_spacing = 1.1
    if justificado:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def titulo_secao(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(5)
    run(p, texto.upper(), tam=10.5, bold=True, cor=VERDE, espaco=20)
    borda_inferior(p)
    return p


def celula(cell, texto, *, tam=8.5, bold=False, italic=False,
           cor=PRETO, centro=False, fundo=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if centro:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, texto, tam=tam, bold=bold, italic=italic, cor=cor)
    if fundo:
        shade(cell, fundo)


# ══════════════════════════════════════════════════════════════════════════
# Carga dos dados da execução
# ══════════════════════════════════════════════════════════════════════════
def carregar():
    scen = pd.read_csv(DATA / 'scen12_best.csv')
    bh   = json.loads((DATA / 'bh_final.json').read_text())
    fc   = json.loads((DATA / 'fc_results_best.json').read_text())

    serie = pd.read_csv(DATA / 'serie_subst.csv')
    anual = serie.groupby('ano')['prec'].sum()
    anual = anual[anual.index < serie['ano'].max() + 1]
    # descartar ano incompleto no fim da série
    cont  = serie.groupby('ano')['mes'].count()
    anual = anual[cont == 12]

    html = (DOCS / 'index.html').read_text(encoding='utf-8')
    meta = {}
    bloco = re.search(r'_meta:\s*\{(.*?)\n  \}', html, re.DOTALL)
    if bloco:
        for k, v in re.findall(r"(\w+):\s*'([^']*)'", bloco.group(1)):
            meta[k] = v

    # trajetória IRI do Niño 3.4 (pico projetado)
    m_iri = re.search(r'iri_nino_mean:\s*\[([^\]]+)\]', html)
    iri_pico = max(float(v) for v in m_iri.group(1).split(',')) if m_iri else None

    # probabilidades CPC/IRI
    m_cpc = re.search(r'cpc:\s*\{\s*en:\[([^\]]+)\]', html)
    m_iri_en = re.search(r'iri:\s*\{\s*en:\[([^\]]+)\]', html)
    cpc_en = [int(v) for v in m_cpc.group(1).split(',')] if m_cpc else []
    iri_en = [int(v) for v in m_iri_en.group(1).split(',')] if m_iri_en else []

    return dict(scen=scen, bh=bh, fc=fc, anual=anual, meta=meta,
                iri_pico=iri_pico, cpc_en=cpc_en, iri_en=iri_en)


def cenario_central(iri_pico):
    """Classifica o cenário de referência pelo pico projetado do Niño 3.4."""
    if iri_pico is None:
        return 'Neutro'
    if iri_pico >= 1.5:  return 'El Nino forte'
    if iri_pico >= 1.0:  return 'El Nino moderado'
    if iri_pico >= 0.5:  return 'El Nino fraco'
    if iri_pico <= -1.5: return 'La Nina forte'
    if iri_pico <= -1.0: return 'La Nina moderada'
    if iri_pico <= -0.5: return 'La Nina fraca'
    return 'Neutro'


NOME_PT = {
    'El Nino forte':    'El Niño forte',
    'El Nino moderado': 'El Niño moderado',
    'El Nino fraco':    'El Niño fraco',
    'Neutro':           'Neutro',
    'La Nina fraca':    'La Niña fraca',
    'La Nina moderada': 'La Niña moderada',
    'La Nina forte':    'La Niña forte',
}


def num(v, dec=0):
    """Formata número no padrão brasileiro."""
    s = f'{v:,.{dec}f}'
    return s.replace(',', '\u00a0').replace('.', ',').replace('\u00a0', '.')


# ══════════════════════════════════════════════════════════════════════════
# Montagem do documento
# ══════════════════════════════════════════════════════════════════════════
def montar(d):
    HOJE   = date.today()
    scen   = d['scen']; bh = d['bh']; fc = d['fc']; meta = d['meta']
    labels = fc[list(fc.keys())[0]]['labels']
    horiz  = f"{MESES[int(labels[0][:2])-1]}/{labels[0][3:]}"
    horiz2 = f"{MESES[int(labels[-1][:2])-1]}/{labels[-1][3:]}"
    horiz_ext = f"{MESES[int(labels[0][:2])-1]}/20{labels[0][3:]} \u2013 {MESES[int(labels[-1][:2])-1]}/20{labels[-1][3:]}"

    central   = cenario_central(d['iri_pico'])
    linha_c   = scen[scen['cenario'] == central].iloc[0]
    media_hist = d['anual'].mean()

    arm_c    = bh['cenarios'][central]['arm']
    arm_clim = bh['clim']['arm']
    n_crit_c    = sum(1 for v in arm_c    if v < ARM_CRITICO)
    n_crit_clim = sum(1 for v in arm_clim if v < ARM_CRITICO)

    def_c    = bh['cenarios'][central]['def_total']
    exc_c    = bh['cenarios'][central]['exc_total']
    def_clim = bh['clim']['def_total']
    exc_clim = bh['clim']['exc_total']

    prec_c = fc[central]['prec']
    # climatologia mensal alinhada ao horizonte
    m0 = int(labels[0][:2])
    clim_jd = [d['anual'].sum()] * 0  # placeholder
    serie = pd.read_csv(DATA / 'serie_subst.csv')
    clim_mensal = serie.groupby('mes')['prec'].mean().round(1).tolist()  # jan..dez
    clim_h = [clim_mensal[(m0 - 1 + i) % 12] for i in range(12)]

    # meses de entrada da estação chuvosa (primeiros com climatologia >= 100 mm)
    rampa = [i for i, v in enumerate(clim_h) if v >= 100][:2]

    # ── documento ───────────────────────────────────────────────────────
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(9.5)

    sec = doc.sections[0]
    sec.top_margin    = Twips(1000)
    sec.bottom_margin = Twips(900)
    sec.left_margin   = Twips(1080)
    sec.right_margin  = Twips(1080)
    LARG = 9746   # largura útil em twips

    # ── cabeçalho ───────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1.5)
    run(p, 'SINOBRAS FLORESTAL  \u00b7  PLANEJAMENTO E OPERAÇÕES',
        tam=7.5, bold=True, cor=VERDE_C, espaco=40)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run(p, 'Relatório Executivo \u2014 Clima Operacional', tam=15, bold=True, cor=VERDE)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    run(p, f'Norte do Tocantins  \u00b7  Horizonte {horiz_ext}  \u00b7  '
           f'Emissão {HOJE.strftime("%d/%m/%Y")}', tam=8.5, cor=CINZA)
    borda_inferior(p, cor=VERDE_HEX, tamanho=12)

    # ── sumário executivo ───────────────────────────────────────────────
    titulo_secao(doc, 'Sumário executivo')

    tipo = ('El Niño' if 'El Nino' in central
            else 'La Niña' if 'La Nina' in central else 'condições neutras')
    intens = central.split()[-1] if central != 'Neutro' else ''

    p = par_corpo(doc)
    if central == 'Neutro':
        run(p, 'O Pacífico equatorial encontra-se em ')
        run(p, 'condições neutras', bold=True)
        run(p, ', sem forçante ENSO dominante para a safra. ')
    else:
        run(p, 'O Pacífico equatorial encontra-se em ')
        run(p, f'{tipo} de intensidade {intens}', bold=True)
        run(p, f', com pico projetado de {num(d["iri_pico"],2).replace(".",",")}\u00a0\u00b0C '
               f'no Niño 3.4. ')
    if d['cpc_en'] and d['iri_en']:
        run(p, f'CPC e IRI convergem em probabilidade de até {max(d["cpc_en"])}% e '
               f'{max(d["iri_en"])}%, respectivamente, para a persistência do evento ao '
               f'longo da estação chuvosa. ')
    seco = linha_c['anom_pct'] < 0
    run(p, f'Para a safra isso se traduz em uma estação chuvosa '
           f'{"mais curta e menos volumosa" if seco else "de volume acima do normal"}, '
           f'com {"déficit hídrico ampliado nas bordas da janela de plantio" if seco else "maior risco de excesso e restrição de tráfego"}.',
        bold=True)

    p = par_corpo(doc)
    run(p, f'O cenário central adotado \u2014 {NOME_PT[central]} \u2014 projeta ')
    run(p, f'{num(linha_c["total12m"])}\u00a0mm', bold=True, cor=SECA if seco else VERDE)
    run(p, ' no horizonte de 12 meses, ')
    run(p, f'{num(abs(linha_c["anom_pct"]),1)}% {"abaixo" if seco else "acima"}',
        bold=True, cor=SECA if seco else VERDE)
    run(p, f' da média histórica de {num(media_hist)}\u00a0mm '
           f'({d["anual"].index.min()}\u2013{d["anual"].index.max()}), com '
           f'{num(linha_c["p_seca"],1)}% de probabilidade de ano seco. '
           f'O núcleo chuvoso de dezembro a março permanece operacionalmente viável, '
           f'com volumes ')
    razao = [prec_c[i] / clim_h[i] for i in range(12)
             if clim_h[i] >= 200 and prec_c[i] > 0]
    if razao:
        d1, d2 = (1 - max(razao)) * 100, (1 - min(razao)) * 100
        lo, hi = sorted([d1, d2])
        run(p, f'{num(abs(lo))}% a {num(abs(hi))}% '
               f'{"menores" if lo >= 0 else "maiores"} que a climatologia.')
    else:
        run(p, 'próximos à climatologia.')

    # ── 1. ENSO ─────────────────────────────────────────────────────────
    titulo_secao(doc, '1. Situação ENSO \u2014 monitoramento')

    p = par_corpo(doc, justificado=False)
    run(p, 'Niño 3.4 semanal: ', bold=True)
    run(p, meta.get('nino34_semana', 'n/d'), bold=True, cor=SECA)
    run(p, '  \u00b7  ')
    run(p, 'ONI: ', bold=True)
    run(p, meta.get('oni_mensal', 'n/d'), bold=True)
    run(p, f'  \u00b7  Pico projetado: ')
    run(p, f'{num(d["iri_pico"],2).replace(".",",")}\u00a0\u00b0C', bold=True)

    p = par_corpo(doc)
    run(p, f'CPC (emissão {meta.get("cpc_emissao","n/d")}): ', bold=True)
    run(p, f'{meta.get("cpc_status","n/d")}. ')
    run(p, f'IRI (Quick Look {meta.get("iri_emissao","n/d")}): ', bold=True)
    if d['iri_en'] and d['cpc_en']:
        div = max(abs(a - b) for a, b in zip(d['iri_en'], d['cpc_en']))
        run(p, f'probabilidade de até {max(d["iri_en"])}% nos trimestres da estação '
               f'chuvosa. A maior divergência entre os dois centros no horizonte é de '
               f'{div} {"ponto percentual" if div == 1 else "pontos percentuais"}. ')
    run(p, 'Divergências acima de 10 pontos percentuais indicam incerteza elevada e '
           'recomendam postura mais conservadora na janela de transição.')

    # ── 2. Projeção ─────────────────────────────────────────────────────
    titulo_secao(doc, '2. Projeção de precipitação por cenário')

    p = par_corpo(doc, depois=4)
    run(p, f'Totais acumulados para o horizonte {horiz_ext}, gerados por SARIMAX '
           f'condicionado a Niño 3.4, TSA e PDO. Déficit e meses críticos referem-se '
           f'ao balanço hídrico com CAD de 100\u00a0mm; meses críticos são aqueles com '
           f'armazenamento de água no solo abaixo de {ARM_CRITICO}\u00a0mm.',
        tam=8.5, cor=CINZA)

    COLS = [2646, 1600, 1500, 1400, 1300, 1300]
    HDR  = ['Cenário ENSO','Total (mm)','Anomalia','P(seca)','Déficit','Meses críticos']
    mostrar = ['El Nino forte','El Nino moderado','El Nino fraco','Neutro','La Nina fraca']

    t = doc.add_table(rows=1, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    bordas_tabela(t); espacamento_celula(t)
    for j, (c, w) in enumerate(zip(t.rows[0].cells, COLS)):
        c.width = Twips(w)
        celula(c, HDR[j], tam=8, bold=True, cor=RGBColor(0xFF,0xFF,0xFF),
               centro=(j > 0), fundo=VERDE_HEX)

    for i, nome in enumerate(mostrar):
        r = scen[scen['cenario'] == nome].iloc[0]
        arm = bh['cenarios'][nome]['arm']
        vals = [
            NOME_PT[nome],
            num(r['total12m']),
            ('+' if r['anom_pct'] > 0 else '\u2212') + num(abs(r['anom_pct']),1) + '%',
            num(r['p_seca'],1) + '%',
            str(bh['cenarios'][nome]['def_total']),
            str(sum(1 for v in arm if v < ARM_CRITICO)),
        ]
        destaque = nome == central
        row = t.add_row()
        for j, (c, w) in enumerate(zip(row.cells, COLS)):
            c.width = Twips(w)
            celula(c, vals[j], tam=8.5, bold=(destaque and j == 0),
                   cor=SECA if destaque else PRETO, centro=(j > 0),
                   fundo=ZEBRA_HEX if i % 2 == 0 else None)

    row = t.add_row()
    clim_vals = [f'Média {d["anual"].index.min()}\u2013{d["anual"].index.max()}',
                 num(media_hist), '\u2014', '\u2014',
                 str(def_clim), str(n_crit_clim)]
    for j, (c, w) in enumerate(zip(row.cells, COLS)):
        c.width = Twips(w)
        celula(c, clim_vals[j], tam=8.5, italic=(j in (0,1,4,5)),
               centro=(j > 0), fundo=CLIM_HEX)

    layout_fixo(t, COLS)

    p = par_corpo(doc, antes=4, depois=0)
    run(p, 'Leitura: ', tam=8.5, bold=True, cor=CINZA)
    delta_def = def_c - def_clim
    delta_exc = round((1 - exc_c / exc_clim) * 100) if exc_clim else 0
    run(p, f'o cenário central eleva o déficit anual em {abs(delta_def)}\u00a0mm '
           f'{"sobre" if delta_def > 0 else "abaixo d"}a climatologia e '
           f'{"adiciona" if n_crit_c > n_crit_clim else "mantém"} '
           f'{abs(n_crit_c - n_crit_clim) if n_crit_c != n_crit_clim else "o mesmo número de"} '
           f'{"mês" if abs(n_crit_c-n_crit_clim)==1 else "meses"} de solo esgotado. '
           f'O excedente hídrico {"cai" if exc_c < exc_clim else "sobe"} de '
           f'{num(exc_clim)}\u00a0mm para {num(exc_c)}\u00a0mm \u2014 '
           f'{"redução" if delta_exc > 0 else "aumento"} de {abs(delta_exc)}% na '
           f'recarga do perfil.', tam=8.5, cor=CINZA)

    # ── 3. Balanço hídrico ──────────────────────────────────────────────
    titulo_secao(doc, '3. Balanço hídrico e implicação de campo')

    seq = []
    atual = 0
    for v in arm_c:
        if v < ARM_CRITICO:
            atual += 1
        else:
            seq.append(atual); atual = 0
    seq.append(atual)
    maior_seq = max(seq)

    p = par_corpo(doc)
    run(p, f'No cenário {NOME_PT[central]}, o solo permanece abaixo de '
           f'{ARM_CRITICO}\u00a0mm de armazenamento por ')
    run(p, f'{n_crit_c} {"mês" if n_crit_c == 1 else "meses"}', bold=True,
        cor=SECA if n_crit_c > n_crit_clim else PRETO)
    run(p, f' do horizonte \u2014 sendo a maior sequência contínua de {maior_seq} '
           f'{"mês" if maior_seq == 1 else "meses"} \u2014, contra {n_crit_clim} na '
           f'climatologia. ')
    if rampa:
        i1 = rampa[0]
        nome_m1 = MESES_EXT[(m0 - 1 + i1) % 12]
        run(p, f'Na entrada da estação chuvosa, {nome_m1} projeta '
               f'{num(prec_c[i1])}\u00a0mm contra {num(clim_h[i1])}\u00a0mm da '
               f'climatologia')
        if len(rampa) > 1:
            i2 = rampa[1]
            nome_m2 = MESES_EXT[(m0 - 1 + i2) % 12]
            run(p, f' e {nome_m2}, {num(prec_c[i2])}\u00a0mm contra '
                   f'{num(clim_h[i2])}\u00a0mm')
        run(p, '.')

    p = par_corpo(doc)
    run(p, 'A consequência operacional é direta: ')
    if seco:
        run(p, 'a recarga do perfil de solo atrasa', bold=True)
        run(p, '. Mudas plantadas na janela de transição enfrentarão maior risco de '
               'estresse hídrico no estabelecimento, justamente na faixa em que '
               'historicamente se antecipa parte do plantio para aliviar o pico '
               'logístico do núcleo chuvoso.')
    else:
        run(p, 'a recarga do perfil ocorre mais cedo e com maior volume', bold=True)
        run(p, '. O risco desloca-se do estresse hídrico para a trafegabilidade e '
               'para o encharcamento em solos de menor drenagem.')

    # ── quebra de página ────────────────────────────────────────────────
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── 4. Recomendação ─────────────────────────────────────────────────
    titulo_secao(doc, '4. Recomendação operacional \u2014 janela de plantio')

    JANELA = [
        ('Dez \u2013 Mar', 'Núcleo',   '1FA05A', 'Plantio pleno \u2014 solo em capacidade de campo'),
        ('Out \u2013 Nov', 'Atenção',  'D18B2C', 'Liberar por talhão, com ONI e chuva acumulada'),
        ('Abr',            'Atenção',  'D18B2C', 'Encerrar plantio; priorizar áreas de menor risco'),
        ('Mai / Set',      'Restrito', 'A8322C', 'Somente com irrigação de salvamento'),
        ('Jun \u2013 Ago', 'Fechado',  '6E6E6E', 'Estação seca estrutural \u2014 sem plantio'),
    ]
    JC = [1800, 1500, 6446]
    t2 = doc.add_table(rows=1, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    bordas_tabela(t2); espacamento_celula(t2)
    for j, (c, w, h) in enumerate(zip(t2.rows[0].cells, JC,
                                      ['Período','Postura','Orientação operacional'])):
        c.width = Twips(w)
        celula(c, h, tam=8, bold=True, cor=RGBColor(0xFF,0xFF,0xFF),
               centro=(j == 1), fundo=VERDE_HEX)
    for i, (per, post, cor, orient) in enumerate(JANELA):
        row = t2.add_row()
        fundo = ZEBRA_HEX if i % 2 == 0 else None
        celula(row.cells[0], per, tam=8.5, bold=True, fundo=fundo)
        celula(row.cells[1], post, tam=8.5, bold=True,
               cor=RGBColor.from_string(cor), centro=True, fundo=fundo)
        celula(row.cells[2], orient, tam=8.5, fundo=fundo)
    layout_fixo(t2, JC)

    p = par_corpo(doc, justificado=False, antes=6, depois=5)
    run(p, f'Ajustes recomendados para a safra:', tam=9.5, bold=True, cor=VERDE)

    p = par_corpo(doc)
    run(p, 'a) Concentrar o plantio no núcleo. ', bold=True)
    run(p, 'Deslocar para o núcleo chuvoso o volume que normalmente seria antecipado '
           'para a janela de transição, aceitando maior pressão logística em troca de '
           'menor risco de replantio.')

    p = par_corpo(doc)
    run(p, 'b) Condicionar a janela de atenção a gatilhos objetivos. ', bold=True)
    run(p, 'Liberar plantio nos meses de atenção apenas por talhão e mediante critérios '
           'de chuva acumulada e armazenamento de solo definidos pela operação \u2014 os '
           'limiares devem ser calibrados com o histórico de pegamento de cada região.')

    p = par_corpo(doc)
    run(p, 'c) Dimensionar contingência de replantio. ', bold=True)
    run(p, 'Provisionar mudas adicionais para as áreas plantadas fora do núcleo, cuja '
           'taxa de perda tende a ser superior à média em anos de forçante ENSO ativa.')

    prox = meta.get('dash_proxima', 'a cada dia 21')
    p = par_corpo(doc)
    run(p, 'd) Revisar mensalmente. ', bold=True)
    run(p, f'O dashboard é reprocessado automaticamente com os dados do mês anterior e '
           f'os índices oceânicos atualizados (próxima atualização: {prox}). ')
    outros = scen[scen['cenario'] != central]
    viz = outros.iloc[(outros['total12m'] - linha_c['total12m']).abs().argsort()].iloc[0]
    run(p, f'Uma reclassificação do evento para {NOME_PT[viz["cenario"]]} altera a '
           f'projeção de {num(linha_c["total12m"])}\u00a0mm para '
           f'{num(viz["total12m"])}\u00a0mm e a probabilidade de ano seco de '
           f'{num(linha_c["p_seca"],1)}% para {num(viz["p_seca"],1)}%, o que '
           f'redefine a postura da janela de transição.')

    # ── rodapé ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run(p, 'Fontes: NOAA/CPC \u00b7 IRI Columbia \u00b7 NASA MERRA-2 \u00b7 '
           'ECMWF ERA5-Land (Open-Meteo) \u00b7 rede de estações Sinobras Florestal. '
           f'Cenário central classificado pelo pico projetado do Niño 3.4 '
           f'({num(d["iri_pico"],2).replace(".",",")}\u00a0\u00b0C). '
           f'Documento gerado automaticamente em {HOJE.strftime("%d/%m/%Y")}.',
        tam=7.5, italic=True, cor=CINZA)
    borda_superior(p)

    return doc


def main():
    print(f"\n{'='*55}")
    print(f"  RELATÓRIO EXECUTIVO — {date.today().strftime('%d/%m/%Y')}")
    print(f"{'='*55}")

    try:
        d = carregar()
    except FileNotFoundError as e:
        print(f"  ❌ Arquivo de dados ausente: {e}")
        return 1

    central = cenario_central(d['iri_pico'])
    print(f"\n  Cenário central : {NOME_PT[central]} (pico Niño 3.4 = {d['iri_pico']})")
    linha = d['scen'][d['scen']['cenario'] == central].iloc[0]
    print(f"  Projeção 12m    : {linha['total12m']:.0f} mm ({linha['anom_pct']:+.1f}%)")
    print(f"  Déficit         : {d['bh']['cenarios'][central]['def_total']} mm "
          f"(clim: {d['bh']['clim']['def_total']} mm)")

    doc = montar(d)
    SAIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SAIDA)
    kb = SAIDA.stat().st_size / 1024
    print(f"\n  ✅ {SAIDA.relative_to(ROOT)} gerado ({kb:.0f} KB)")
    print(f"{'='*55}\n")
    return 0


if __name__ == '__main__':
    sys.exit(main())
